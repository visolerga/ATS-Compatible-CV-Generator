from app.models.cv_model import CVModel
from app.models.user_model import User
from fpdf import FPDF
from docx import Document
from jinja2 import Environment, FileSystemLoader
import os

class CVService:
    """Servicio para la generación de CVs en diferentes formatos."""

    @staticmethod
    def create_cv(user_id, personal_info, professional_summary, work_experience,
                  education, skills, certifications, projects, references=None):
        """Crea un nuevo CV para un usuario."""
        try:
            cv = CVModel.create(
                user_id=user_id,
                personal_info=personal_info,
                professional_summary=professional_summary,
                work_experience=work_experience,
                education=education,
                skills=skills,
                certifications=certifications,
                projects=projects,
                references=references
            )
            return cv
        except Exception as e:
            print(f"Error al crear el CV: {e}")
            return None

    @staticmethod
    def generate_pdf(cv_id, file_path):
        """Genera un CV en formato PDF."""
        try:
            cv = CVModel.get(CVModel.id == cv_id)
            pdf = FPDF()
            pdf.add_page()

            pdf.set_font("Times", "B", 16)
            pdf.cell(0, 10, "Curriculum Vitae", 0, 1, 'C')

            pdf.set_font("Times", "", 12)
            # Agregar información personal
            personal_info = cv.personal_info
            pdf.multi_cell(0, 10, f"{personal_info['name']}\n{personal_info['address']}\n{personal_info['phone']}\n{personal_info['email']}")
            pdf.cell(0, 10, "", 0, 1)  # Espacio

            pdf.set_font("Georgia", "I", 12)
            pdf.cell(0, 10, "Resumen Profesional", 0, 1)
            pdf.set_font("Times", "", 12)
            pdf.multi_cell(0, 10, cv.professional_summary)

            # Agregar Experiencia Laboral
            pdf.set_font("Georgia", "I", 12)
            pdf.cell(0, 10, "Experiencia Laboral", 0, 1)
            pdf.set_font("Times", "", 12)
            for job in cv.work_experience:
                pdf.cell(0, 10, f"{job['job_title']} en {job['company']} ({job['duration']})", 0, 1)
                pdf.multi_cell(0, 10, "\n".join(job['responsibilities']))
                pdf.cell(0, 10, "", 0, 1)  # Espacio

            pdf.output(file_path)
        except Exception as e:
            print(f"Error al generar el PDF: {e}")

    @staticmethod
    def generate_word(cv_id, file_path):
        """Genera un CV en formato Word (DOCX)."""
        try:
            cv = CVModel.get(CVModel.id == cv_id)
            doc = Document()

            doc.add_heading("Curriculum Vitae", level=1)

            # Agregar información personal
            personal_info = cv.personal_info
            doc.add_heading("Información Personal", level=2)
            doc.add_paragraph(f"{personal_info['name']}\n{personal_info['address']}\n{personal_info['phone']}\n{personal_info['email']}")

            doc.add_heading("Resumen Profesional", level=2)
            doc.add_paragraph(cv.professional_summary)

            # Agregar Experiencia Laboral
            doc.add_heading("Experiencia Laboral", level=2)
            for job in cv.work_experience:
                doc.add_paragraph(f"{job['job_title']} en {job['company']} ({job['duration']})")
                for responsibility in job['responsibilities']:
                    doc.add_paragraph(f" - {responsibility}")

            doc.save(file_path)
        except Exception as e:
            print(f"Error al generar el archivo Word: {e}")

    @staticmethod
    def generate_html(cv_id, file_path):
        """Genera un CV en formato HTML."""
        try:
            cv = CVModel.get(CVModel.id == cv_id)

            # Cargar la plantilla HTML
            env = Environment(loader=FileSystemLoader('templates'))
            template = env.get_template('cv_template.html')

            # Renderizar la plantilla con datos del CV
            html_output = template.render(
                personal_info=cv.personal_info,
                professional_summary=cv.professional_summary,
                work_experience=cv.work_experience,
                education=cv.education,
                skills=cv.skills,
                certifications=cv.certifications,
                projects=cv.projects,
                references=cv.references
            )

            with open(file_path, 'w') as f:
                f.write(html_output)
        except Exception as e:
            print(f"Error al generar el archivo HTML: {e}")
